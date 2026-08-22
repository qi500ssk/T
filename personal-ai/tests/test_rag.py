import io
import json
import math
import zipfile

import pytest
from docx import Document as WordDocument
from pypdf import PdfWriter

from core.rag.embedding import MockEmbeddingProvider
from core.chat.gateway import MockProvider, StreamChunk
from core.rag.chunking import split_into_chunks
from core.rag.parsers import (
    ParsedBlock,
    ResourceLimitError,
    parse_docx,
    parse_markdown,
    parse_pdf,
    parse_plain_text,
)
from core.rag.retrieval import should_retrieve_knowledge, tokenize_for_bm25
from infrastructure.config import settings


def parse_sse(text: str) -> list[tuple[str, dict]]:
    events = []
    for block in text.split("\n\n"):
        if not block.strip():
            continue
        event = ""
        data = "{}"
        for line in block.splitlines():
            if line.startswith("event:"):
                event = line[6:].strip()
            elif line.startswith("data:"):
                data = line[5:].strip()
        events.append((event, json.loads(data)))
    return events


def test_mock_embedding_is_deterministic_normalized_and_dimensioned():
    provider = MockEmbeddingProvider(64)
    first = provider.embed_query("火星基地")
    second = provider.embed_documents(["火星基地"])[0]
    assert first == second
    assert len(first) == 64
    assert math.isclose(sum(value * value for value in first), 1.0)


def test_rag_query_gate_skips_non_knowledge_requests():
    for query in ["你好", "计算出122+22", "请算一下 (12+3)*4", "现在几点？", "查询当前时间和地点还有天气", "帮我写入笔记：P6完成"]:
        assert should_retrieve_knowledge(query) is False
    for query in ["火星计划代号是什么？", "根据资料总结项目", "简历中的工作经历是什么？"]:
        assert should_retrieve_knowledge(query) is True
    assert "+" not in tokenize_for_bm25("Agent+RAG 122+22")


def test_text_markdown_and_docx_parsers_preserve_structure(tmp_path):
    text_path = tmp_path / "notes.txt"
    text_path.write_text("第一段。\n\n第二段。", encoding="utf-8")
    plain = parse_plain_text(text_path, settings)
    assert [block.content for block in plain.blocks] == ["第一段。", "第二段。"]
    assert plain.blocks[1].char_start is not None

    md_path = tmp_path / "guide.md"
    md_path.write_text("# 安装\n说明。\n## Windows\n步骤。", encoding="utf-8")
    markdown = parse_markdown(md_path, settings)
    assert [block.section for block in markdown.blocks] == ["安装", "安装 > Windows"]

    docx_path = tmp_path / "guide.docx"
    document = WordDocument()
    document.add_heading("部署", level=1)
    document.add_heading("Linux", level=2)
    document.add_paragraph("执行安装命令。")
    document.save(docx_path)
    docx = parse_docx(docx_path, settings)
    assert docx.blocks[0].section == "部署 > Linux"


def test_blank_pdf_needs_ocr(tmp_path):
    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    with path.open("wb") as output:
        writer.write(output)
    result = parse_pdf(path, settings)
    assert result.needs_ocr is True
    assert result.blocks == []


def test_docx_zip_resource_limit(tmp_path, monkeypatch):
    path = tmp_path / "large.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", "x")
        archive.writestr("word/document.xml", "x" * 100)
    monkeypatch.setattr(settings, "docx_max_uncompressed_bytes", 50)
    with pytest.raises(ResourceLimitError):
        parse_docx(path, settings)


def test_chunking_obeys_character_and_token_limits(monkeypatch):
    monkeypatch.setattr(settings, "rag_chunk_target_chars", 20)
    monkeypatch.setattr(settings, "rag_chunk_max_chars", 30)
    monkeypatch.setattr(settings, "rag_chunk_max_tokens", 15)
    monkeypatch.setattr(settings, "rag_chunk_overlap_sentences", 1)
    blocks = [ParsedBlock("章节", "第一句话很完整。第二句话也完整。第三句话仍然完整。", char_start=10)]
    chunks = split_into_chunks(blocks, lambda text: max(1, len(text) // 2), settings)
    assert chunks
    assert all(len(chunk.content) <= 30 for chunk in chunks)
    assert all(len(chunk.content) // 2 <= 15 for chunk in chunks)
    assert all(chunk.section == "章节" for chunk in chunks)
    assert all(chunk.char_start is not None and chunk.char_end is not None for chunk in chunks)


def test_upload_search_chat_citations_and_delete(client):
    content = "# 火星计划\n火星计划的内部代号是晨星七号，负责人是林岚。\n"
    upload = client.post(
        "/api/files",
        files={"file": ("火星计划.md", content.encode("utf-8"), "text/markdown")},
    )
    assert upload.status_code == 201
    document = upload.json()
    assert document["status"] == "indexed"
    assert document["chunk_count"] >= 1

    rows = client.get("/api/documents").json()
    assert rows[0]["id"] == document["id"]
    detail = client.get(f"/api/documents/{document['id']}").json()
    assert detail["chunks"][0]["section"] == "火星计划"
    original = client.get(f"/api/documents/{document['id']}/content")
    assert original.status_code == 200
    assert "晨星七号" in original.content.decode("utf-8")

    search = client.get("/api/search", params={"q": "晨星七号", "limit": 5})
    assert search.status_code == 200
    assert search.json()[0]["document_id"] == document["id"]

    conversation = client.post("/api/conversations", json={}).json()
    response = client.post(
        "/api/chat",
        json={"conversation_id": conversation["id"], "message": "火星计划代号是什么？"},
    )
    events = parse_sse(response.text)
    types = [event for event, _ in events]
    assert "rag.retrieved" in types
    assert types.index("message.delta") < types.index("rag.retrieved")
    assert types.index("rag.retrieved") < types.index("message.completed")
    sources = next(data["sources"] for event, data in events if event == "rag.retrieved")
    assert sources[0]["citation_id"] == "c1"
    assert sources[0]["char_start"] is not None

    messages = client.get(f"/api/conversations/{conversation['id']}/messages").json()
    assistant = messages[-1]
    assert "[c1]" in assistant["content"]
    assert assistant["citations"][0]["document_id"] == document["id"]

    deleted = client.delete(f"/api/documents/{document['id']}")
    assert deleted.status_code == 200
    assert client.get("/api/documents").json() == []
    assert client.get(f"/api/documents/{document['id']}/content").status_code == 404


def test_calculation_bypasses_rag_even_when_documents_exist(client, monkeypatch):
    monkeypatch.setattr(settings, "memory_enabled", False)
    upload = client.post(
        "/api/files",
        files={"file": ("无关资料.md", "# 项目\n项目负责人是林岚。".encode("utf-8"), "text/markdown")},
    )
    assert upload.status_code == 201
    conversation = client.post("/api/conversations", json={}).json()
    response = client.post(
        "/api/chat",
        json={"conversation_id": conversation["id"], "message": "计算出122+22"},
    )
    events = parse_sse(response.text)
    assert "rag.retrieved" not in [event for event, _ in events]
    messages = client.get(f"/api/conversations/{conversation['id']}/messages").json()
    assert "144" in messages[-1]["content"]
    assert messages[-1]["citations"] == []


def test_uncited_retrieval_candidates_are_not_exposed_as_sources(client, monkeypatch):
    class NoCitationProvider(MockProvider):
        async def stream(self, messages, temperature=0.7, tools=None):
            yield StreamChunk(text="这是未引用资料的回答。")
            yield StreamChunk(
                finish_reason="stop",
                usage={"prompt_tokens": 1, "completion_tokens": 1},
            )

    monkeypatch.setattr(settings, "memory_enabled", False)
    client.app.state.provider = NoCitationProvider(delay=0)
    upload = client.post(
        "/api/files",
        files={
            "file": (
                "火星计划.md",
                "# 火星计划\n内部代号是晨星七号。".encode("utf-8"),
                "text/markdown",
            )
        },
    )
    assert upload.status_code == 201
    conversation = client.post("/api/conversations", json={}).json()
    response = client.post(
        "/api/chat",
        json={"conversation_id": conversation["id"], "message": "火星计划代号是什么？"},
    )
    assert "rag.retrieved" not in [event for event, _ in parse_sse(response.text)]
    messages = client.get(f"/api/conversations/{conversation['id']}/messages").json()
    assert messages[-1]["citations"] == []


def test_duplicate_and_disguised_uploads_are_rejected(client):
    payload = {"file": ("same.txt", "唯一内容".encode("utf-8"), "text/plain")}
    first = client.post("/api/files", files=payload)
    assert first.status_code == 201
    duplicate = client.post("/api/files", files=payload)
    assert duplicate.status_code == 409
    assert duplicate.json()["detail"]["document_id"] == first.json()["id"]

    disguised = client.post(
        "/api/files",
        files={"file": ("fake.pdf", b"not a pdf", "application/pdf")},
    )
    assert disguised.status_code == 415
    assert len(client.get("/api/documents").json()) == 1


def test_upload_size_limit_does_not_create_document(client, monkeypatch):
    monkeypatch.setattr(settings, "file_max_bytes", 4)
    response = client.post(
        "/api/files",
        files={"file": ("too-large.txt", b"12345", "text/plain")},
    )
    assert response.status_code == 413
    assert client.get("/api/documents").json() == []
