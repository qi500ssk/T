import asyncio
import json
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from pptx import Presentation

from core.files.artifacts import resolve_artifact
from mcp_servers.document_skills.generation import (
    create_docx,
    create_pdf,
    create_pptx,
    create_xlsx,
)
from core.capabilities.mcp import McpClient, McpServerConfig
from infrastructure.config import settings


def _artifact_id(result: str) -> str:
    assert result.startswith("ARTIFACT_JSON:")
    return json.loads(result.splitlines()[0].removeprefix("ARTIFACT_JSON:"))["id"]


def test_generate_all_document_formats_and_validate_structure(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "artifacts_dir", str(tmp_path / "artifacts"))

    docx_id = _artifact_id(create_docx("测试报告", "# 摘要\n正文内容。\n- 要点一", "测试报告.docx"))
    docx_record = resolve_artifact(docx_id)
    document = Document(docx_record.path)
    assert document.paragraphs[0].text == "测试报告"
    assert any(item.text == "摘要" for item in document.paragraphs)

    pdf_id = _artifact_id(create_pdf("测试 PDF", "# 摘要\n支持中文正文。", "测试.pdf"))
    pdf_record = resolve_artifact(pdf_id)
    assert len(PdfReader(pdf_record.path).pages) == 1

    pptx_id = _artifact_id(create_pptx(
        "测试演示",
        [
            {"title": "目标", "bullets": ["明确范围", "完成验证"]},
            {"title": "结果", "bullets": ["文件可下载"]},
        ],
        "测试演示.pptx",
    ))
    presentation = Presentation(resolve_artifact(pptx_id).path)
    assert len(presentation.slides) == 3

    xlsx_id = _artifact_id(create_xlsx(
        "测试工作簿",
        [{"name": "数据", "rows": [["项目", "数量", "合计"], ["A", 2, "=B2*3"]]}],
        "测试工作簿.xlsx",
    ))
    workbook = load_workbook(resolve_artifact(xlsx_id).path, data_only=False)
    assert workbook["数据"]["B2"].value == 2
    assert workbook["数据"]["C2"].value == "=B2*3"


def test_artifact_download_api(client, tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "artifacts_dir", str(tmp_path / "artifacts"))
    artifact_id = _artifact_id(create_docx("可下载", "正文。", "download.docx"))
    response = client.get(f"/api/artifacts/{artifact_id}")
    assert response.status_code == 200
    assert response.content.startswith(b"PK")
    assert "download.docx" in response.headers["content-disposition"]
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert client.get("/api/artifacts/not-an-artifact-id").status_code == 404


@pytest.mark.asyncio
async def test_document_mcp_discovers_tools_and_generates(tmp_path):
    config = McpServerConfig(
        name="document-test",
        command="python",
        args=("-m", "mcp_servers.document_skills.server"),
        env={"ARTIFACTS_DIR": str(tmp_path / "artifacts")},
        enabled=True,
        default_risk_level="medium",
    )
    client = McpClient(config, cwd=Path.cwd())
    try:
        await asyncio.wait_for(client.connect(), 10)
        names = {tool.name for tool in await client.list_tools()}
        assert {"create_docx", "append_docx", "create_pdf", "create_pptx", "create_xlsx"} <= names
        result = await asyncio.wait_for(client.call_tool(
            "create_xlsx",
            {
                "title": "MCP Test",
                "sheets_json": json.dumps([{"name": "Sheet1", "rows": [["Value"], [1]]}]),
                "filename": "mcp-test.xlsx",
            },
        ), 15)
        payload = json.loads(result.splitlines()[0].removeprefix("ARTIFACT_JSON:"))
        assert (tmp_path / "artifacts" / payload["id"] / "artifact.xlsx").is_file()
    finally:
        await asyncio.wait_for(client.close(), 10)
